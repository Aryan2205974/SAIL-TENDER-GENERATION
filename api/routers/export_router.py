import os
import io
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse, FileResponse
from sqlalchemy.orm import Session
from database import get_db
import models
from auth import get_current_user

router = APIRouter(prefix="/api/export", tags=["Export"])


def _get_tender_content(tender_id: int, db: Session) -> tuple[models.Tender, str]:
    tender = db.query(models.Tender).filter(models.Tender.id == tender_id).first()
    if not tender:
        raise HTTPException(404, "Tender not found")

    if tender.draft_content:
        content = tender.draft_content
    else:
        parts = []
        for sec in sorted(tender.sections, key=lambda s: s.order_index):
            parts.append(sec.content or f"## {sec.section_name}\n\n*Content pending*\n")
        content = "\n\n".join(parts) if parts else "No content generated yet."

    return tender, content


@router.get("/pdf/{tender_id}")
def export_pdf(
    tender_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    tender, content = _get_tender_content(tender_id, db)

    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import mm

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=25*mm, rightMargin=25*mm,
                                topMargin=25*mm, bottomMargin=25*mm)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle('TenderTitle', parent=styles['Heading1'], fontSize=16,
                                      spaceAfter=12)
        heading_style = ParagraphStyle('TenderH2', parent=styles['Heading2'], fontSize=13,
                                        spaceAfter=8)
        body_style = ParagraphStyle('TenderBody', parent=styles['Normal'], fontSize=10,
                                     leading=14, spaceAfter=6)

        story = []
        story.append(Paragraph(tender.title, title_style))
        story.append(Paragraph(f"Tender ID: {tender.tender_id}", body_style))
        story.append(Spacer(1, 12))

        for line in content.split("\n"):
            line = line.strip()
            if not line:
                story.append(Spacer(1, 6))
            elif line.startswith("## "):
                story.append(Paragraph(line[3:], heading_style))
            elif line.startswith("### "):
                story.append(Paragraph(line[4:], ParagraphStyle('H3', parent=styles['Heading3'],
                                                                  fontSize=11, spaceAfter=6)))
            elif line.startswith("**") and line.endswith("**"):
                story.append(Paragraph(f"<b>{line[2:-2]}</b>", body_style))
            else:
                safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(safe, body_style))

        doc.build(story)
        buf.seek(0)

        db.add(models.AuditLog(
            tender_id=tender.id,
            action="PDF exported",
            performed_by=current_user.name,
        ))
        db.commit()

        return StreamingResponse(
            buf,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={tender.tender_id.replace('/', '_')}.pdf"}
        )
    except ImportError:
        raise HTTPException(500, "reportlab not installed")


@router.get("/docx/{tender_id}")
def export_docx(
    tender_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    tender, content = _get_tender_content(tender_id, db)

    try:
        from docx import Document
        from docx.shared import Pt, Inches

        doc = Document()
        doc.add_heading(tender.title, level=0)
        doc.add_paragraph(f"Tender ID: {tender.tender_id}")
        doc.add_paragraph("")

        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=1)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=2)
            elif line.startswith("**") and line.endswith("**"):
                p = doc.add_paragraph()
                run = p.add_run(line[2:-2])
                run.bold = True
            else:
                doc.add_paragraph(line)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        db.add(models.AuditLog(
            tender_id=tender.id,
            action="DOCX exported",
            performed_by=current_user.name,
        ))
        db.commit()

        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={tender.tender_id.replace('/', '_')}.docx"}
        )
    except ImportError:
        raise HTTPException(500, "python-docx not installed")


@router.get("/html/{tender_id}")
def export_html(
    tender_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    tender, content = _get_tender_content(tender_id, db)

    try:
        import markdown
        html_content = markdown.markdown(content)
    except ImportError:
        html_content = f"<pre>{content}</pre>"

    html_doc = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>{tender.title}</title>
    <style>
        body {{ font-family: 'Segoe UI', Tahoma, sans-serif; max-width: 900px; margin: 40px auto; padding: 20px; }}
        h1 {{ color: #1a365d; border-bottom: 2px solid #2b6cb0; padding-bottom: 10px; }}
        h2 {{ color: #2b6cb0; margin-top: 30px; }}
        h3 {{ color: #4a5568; }}
    </style>
</head>
<body>
    <h1>{tender.title}</h1>
    <p><strong>Tender ID:</strong> {tender.tender_id}</p>
    {html_content}
</body>
</html>"""

    db.add(models.AuditLog(
        tender_id=tender.id,
        action="HTML exported",
        performed_by=current_user.name,
    ))
    db.commit()

    buf = io.BytesIO(html_doc.encode("utf-8"))
    return StreamingResponse(
        buf,
        media_type="text/html",
        headers={"Content-Disposition": f"attachment; filename={tender.tender_id.replace('/', '_')}.html"}
    )
